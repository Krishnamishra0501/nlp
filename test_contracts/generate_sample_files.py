import os
import fitz # PyMuPDF
import docx

TEST_DIR = os.path.dirname(os.path.abspath(__file__))

def create_sample_files():
    print(f"Creating sample contract files in {TEST_DIR}...")

    # 1. Employment TXT
    emp_text = """EMPLOYMENT AGREEMENT

THIS AGREEMENT is entered into as of January 15, 2026, by and between Apex Tech Inc. ("Employer") and Jane Smith ("Employee").

1. POSITION AND RESPONSIBILITIES
Employer hereby hires Employee as Senior Software Architect. Employee agrees to perform duties diligent and full-time.

2. COMPENSATION AND SALARY
Employer shall pay Employee a base annual salary of $160,000, payable bi-weekly. Employee shall be eligible for health insurance benefits and 401(k) matching.

3. TERMINATION
Either party may terminate employment with 30 days written notice. Employment is at-will under California law.

IN WITNESS WHEREOF, the parties hereto have executed this Employment Agreement.
"""
    txt_path = os.path.join(TEST_DIR, "sample_employment.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(emp_text)
    print(f"  Created: {txt_path}")

    # 2. NDA DOCX
    nda_text_lines = [
        "MUTUAL NON-DISCLOSURE AGREEMENT",
        "This Mutual Non-Disclosure Agreement (\"Agreement\") is made by and between Party A and Party B.",
        "1. CONFIDENTIAL INFORMATION",
        "Confidential Information includes all non-public technical data, source code, trade secrets, business plans, and customer lists.",
        "2. OBLIGATIONS",
        "The Receiving Party shall hold all Confidential Information in strict confidence and shall not disclose it to third parties.",
        "3. TERM",
        "This obligation shall survive for a period of five (5) years from the date of disclosure."
    ]
    docx_path = os.path.join(TEST_DIR, "sample_nda.docx")
    doc = docx.Document()
    doc.add_heading("MUTUAL NON-DISCLOSURE AGREEMENT", level=1)
    for line in nda_text_lines[1:]:
        doc.add_paragraph(line)
    doc.save(docx_path)
    print(f"  Created: {docx_path}")

    # 3. Software License PDF
    pdf_text = """SOFTWARE LICENSE AND SAAS SERVICE AGREEMENT

THIS SAAS AGREEMENT is made effective as of February 1, 2026, by Licensor Inc. and Customer Corp.

SECTION 1. SAAS LICENSE GRANT
Licensor grants Customer a non-exclusive, non-transferable right to access the Cloud Software platform during the Subscription Term.

SECTION 2. USE RESTRICTIONS
Customer shall not reverse engineer, decompile, or disassemble the Software source code.

SECTION 3. SERVICE LEVEL AGREEMENT (SLA)
Licensor guarantees 99.9% uptime for cloud services.

IN WITNESS WHEREOF, the parties have executed this Agreement.
"""
    pdf_path = os.path.join(TEST_DIR, "sample_software_license.pdf")
    pdf_doc = fitz.open()
    page = pdf_doc.new_page()
    page.insert_text((50, 50), pdf_text, fontsize=11)
    pdf_doc.save(pdf_path)
    pdf_doc.close()
    print(f"  Created: {pdf_path}")

if __name__ == "__main__":
    create_sample_files()
